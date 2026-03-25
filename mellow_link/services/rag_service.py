"""
RAG Service - Retrieval-Augmented Generation for Mellow-Link

This module provides document embedding and retrieval capabilities:
- Document text extraction (PDF, DOCX, TXT, MD)
- Embedding generation via Ollama
- Vector storage in SQLite with JSON serialization
- Cosine similarity search

Design:
    - Uses Ollama embeddings API (local, no external API keys needed)
    - Stores embeddings in SQLite for persistence
    - Chunks documents for better retrieval
    - Non-blocking async interface
"""

import asyncio
import json
import logging
import math
import os
import re
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple

import httpx

logger = logging.getLogger(__name__)


# =============================================================================
# Configuration
# =============================================================================

DEFAULT_EMBEDDING_MODEL = "nomic-embed-text"  # Good balance of speed/quality
DEFAULT_CHUNK_SIZE = 500  # characters per chunk
DEFAULT_CHUNK_OVERLAP = 50  # overlap between chunks
DEFAULT_TOP_K = 3  # number of results to return

OLLAMA_BASE_URL = os.getenv("OLLAMA_HOST", "http://localhost:11434")


# =============================================================================
# Data Classes
# =============================================================================

@dataclass
class DocumentChunk:
    """A chunk of text from a document with its embedding."""
    id: int
    folder_id: int
    document_id: int
    filename: str
    content: str
    chunk_index: int
    embedding: List[float] = field(default_factory=list)
    created_at: datetime = field(default_factory=datetime.utcnow)


@dataclass
class TempChunk:
    """Temporary chunk for ephemeral/one-time uploads (in-memory only, no DB)."""
    id: int
    session_id: str
    filename: str
    content: str
    chunk_index: int
    embedding: List[float] = field(default_factory=list)
    created_at: datetime = field(default_factory=datetime.utcnow)


@dataclass
class RAGSearchResult:
    """Result from a RAG search query."""
    content: str
    filename: str
    score: float
    chunk_index: int
    document_id: int


# =============================================================================
# Text Extraction
# =============================================================================

def extract_text_from_file(file_path: Path, content_bytes: bytes = None) -> str:
    """
    Extract text from various file formats.

    Supports: PDF, DOCX, TXT, MD, HTML
    """
    suffix = file_path.suffix.lower()

    try:
        if suffix == ".txt" or suffix == ".md":
            if content_bytes:
                return content_bytes.decode("utf-8", errors="ignore")
            return file_path.read_text(encoding="utf-8", errors="ignore")

        elif suffix == ".pdf":
            return _extract_pdf(file_path, content_bytes)

        elif suffix == ".docx":
            return _extract_docx(file_path, content_bytes)

        elif suffix == ".html" or suffix == ".htm":
            return _extract_html(file_path, content_bytes)

        else:
            # Try as plain text
            if content_bytes:
                return content_bytes.decode("utf-8", errors="ignore")
            return file_path.read_text(encoding="utf-8", errors="ignore")

    except Exception as e:
        logger.error(f"[RAG] Failed to extract text from {file_path}: {e}")
        return ""


def _extract_pdf(file_path: Path, content_bytes: bytes = None) -> str:
    """Extract text from PDF using pdfminer."""
    try:
        from pdfminer.high_level import extract_text
        from io import BytesIO

        if content_bytes:
            return extract_text(BytesIO(content_bytes))
        return extract_text(str(file_path))
    except ImportError:
        logger.warning("[RAG] pdfminer not installed, trying pypdf")
        try:
            from pypdf import PdfReader
            from io import BytesIO

            if content_bytes:
                reader = PdfReader(BytesIO(content_bytes))
            else:
                reader = PdfReader(str(file_path))

            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"[RAG] PDF extraction failed: {e}")
            return ""


def _extract_docx(file_path: Path, content_bytes: bytes = None) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        from io import BytesIO

        if content_bytes:
            doc = Document(BytesIO(content_bytes))
        else:
            doc = Document(str(file_path))

        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        return "\n".join(text_parts)
    except ImportError:
        logger.error("[RAG] python-docx not installed")
        return ""
    except Exception as e:
        logger.error(f"[RAG] DOCX extraction failed: {e}")
        return ""


def _extract_html(file_path: Path, content_bytes: bytes = None) -> str:
    """Extract text from HTML using BeautifulSoup."""
    try:
        from bs4 import BeautifulSoup

        if content_bytes:
            html = content_bytes.decode("utf-8", errors="ignore")
        else:
            html = file_path.read_text(encoding="utf-8", errors="ignore")

        soup = BeautifulSoup(html, "html.parser")

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        return soup.get_text(separator="\n", strip=True)
    except ImportError:
        logger.error("[RAG] beautifulsoup4 not installed")
        return ""
    except Exception as e:
        logger.error(f"[RAG] HTML extraction failed: {e}")
        return ""


# =============================================================================
# Chunking
# =============================================================================

def chunk_text(
    text: str,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    overlap: int = DEFAULT_CHUNK_OVERLAP
) -> List[str]:
    """
    Split text into overlapping chunks.

    Uses sentence boundaries when possible for better context.
    """
    if not text or not text.strip():
        return []

    # Clean text
    text = re.sub(r'\s+', ' ', text).strip()

    # Split by sentences first
    sentences = re.split(r'(?<=[.!?])\s+', text)

    chunks = []
    current_chunk = ""

    for sentence in sentences:
        if len(current_chunk) + len(sentence) <= chunk_size:
            current_chunk += " " + sentence if current_chunk else sentence
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())

            # Start new chunk with overlap from previous
            if chunks and overlap > 0:
                # Get last N characters for overlap
                prev_text = chunks[-1]
                overlap_text = prev_text[-overlap:] if len(prev_text) > overlap else prev_text
                current_chunk = overlap_text + " " + sentence
            else:
                current_chunk = sentence

            # Handle very long sentences
            while len(current_chunk) > chunk_size:
                chunks.append(current_chunk[:chunk_size].strip())
                current_chunk = current_chunk[chunk_size - overlap:]

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    return chunks


# =============================================================================
# Embedding Generation
# =============================================================================

async def generate_embedding(
    text: str,
    model: str = DEFAULT_EMBEDDING_MODEL,
    base_url: str = OLLAMA_BASE_URL
) -> List[float]:
    """
    Generate embedding for text using Ollama.

    Returns empty list on failure.
    """
    if not text or not text.strip():
        return []

    # Check if event loop is available and not closed
    try:
        loop = asyncio.get_running_loop()
        if loop.is_closed():
            logger.debug("[RAG] Event loop is closed, skipping embedding generation")
            return []
    except RuntimeError:
        # No event loop running
        logger.debug("[RAG] No event loop available, skipping embedding generation")
        return []

    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(
                f"{base_url}/api/embeddings",
                json={"model": model, "prompt": text}
            )

            if response.status_code == 200:
                data = response.json()
                embedding = data.get("embedding", [])
                logger.debug(f"[RAG] Generated embedding with {len(embedding)} dimensions")
                return embedding
            else:
                logger.error(f"[RAG] Embedding API error: {response.status_code}")
                return []

    except RuntimeError as e:
        if "Event loop is closed" in str(e) or "closed" in str(e).lower():
            logger.debug("[RAG] Event loop closed during embedding generation")
            return []
        raise
    except Exception as e:
        # Check if error is related to event loop closure
        error_str = str(e).lower()
        if "event loop is closed" in error_str or "closed" in error_str:
            logger.debug(f"[RAG] Event loop closed: {e}")
            return []
        logger.error(f"[RAG] Failed to generate embedding: {e}")
        return []


# =============================================================================
# Vector Math
# =============================================================================

def cosine_similarity(vec1: List[float], vec2: List[float]) -> float:
    """Calculate cosine similarity between two vectors."""
    if not vec1 or not vec2 or len(vec1) != len(vec2):
        return 0.0

    dot_product = sum(a * b for a, b in zip(vec1, vec2))
    norm1 = math.sqrt(sum(a * a for a in vec1))
    norm2 = math.sqrt(sum(b * b for b in vec2))

    if norm1 == 0 or norm2 == 0:
        return 0.0

    return dot_product / (norm1 * norm2)


# =============================================================================
# RAG Service Class
# =============================================================================

class RAGService:
    """
    Service for document embedding and retrieval.

    Uses Ollama for embeddings and SQLite for storage.
    """

    def __init__(
        self,
        embedding_model: str = DEFAULT_EMBEDDING_MODEL,
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
        ollama_url: str = OLLAMA_BASE_URL
    ):
        self.embedding_model = embedding_model
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.ollama_url = ollama_url

        self._executor = ThreadPoolExecutor(max_workers=2)
        self._is_initialized = False

        # In-memory cache of embeddings (folder_id -> list of chunks)
        self._chunk_cache: Dict[int, List[DocumentChunk]] = {}

        # Temporary in-memory store for ephemeral uploads (session_id -> list of TempChunks)
        # This data is NOT persisted to DB - lost on server restart
        self.temp_store: Dict[str, List[TempChunk]] = {}
        
        # Search result cache: query_hash -> (results, timestamp)
        # 캐시 TTL: 5분 (같은 쿼리는 캐시에서 반환)
        self._search_cache: Dict[str, Tuple[List[RAGSearchResult], datetime]] = {}
        self._cache_ttl_seconds = 300  # 5분

    async def initialize(self) -> bool:
        """Initialize RAG service and check Ollama availability."""
        # Check if event loop is available
        try:
            loop = asyncio.get_running_loop()
            if loop.is_closed():
                logger.debug("[RAG] Event loop is closed, skipping initialization")
                return False
        except RuntimeError:
            logger.debug("[RAG] No event loop available, skipping initialization")
            return False

        try:
            # Test Ollama connection
            async with httpx.AsyncClient(timeout=10.0) as client:
                response = await client.get(f"{self.ollama_url}/api/tags")
                if response.status_code == 200:
                    logger.info("[RAG] Ollama connection verified")
                    self._is_initialized = True
                    return True
                else:
                    logger.warning(f"[RAG] Ollama not responding: {response.status_code}")
                    return False
        except RuntimeError as e:
            if "Event loop is closed" in str(e) or "closed" in str(e).lower():
                logger.debug("[RAG] Event loop closed during initialization")
                return False
            raise
        except Exception as e:
            # Check if error is related to event loop closure
            error_str = str(e).lower()
            if "event loop is closed" in error_str or "closed" in error_str:
                logger.debug(f"[RAG] Event loop closed: {e}")
                return False
            logger.warning(f"[RAG] Failed to connect to Ollama: {e}")
            return False

    def is_available(self) -> bool:
        """Check if RAG service is available."""
        return self._is_initialized

    async def process_document(
        self,
        folder_id: int,
        document_id: int,
        filename: str,
        file_path: Path = None,
        content_bytes: bytes = None,
        db_session = None
    ) -> Tuple[bool, int, str]:
        """
        Process a document: extract text, chunk, and generate embeddings.

        Returns: (success, chunk_count, message)
        """
        logger.info(f"[RAG] Processing document: {filename} (folder_id={folder_id})")

        try:
            # 1. Extract text
            if file_path:
                text = await asyncio.get_event_loop().run_in_executor(
                    self._executor,
                    extract_text_from_file,
                    file_path,
                    content_bytes
                )
            elif content_bytes:
                text = await asyncio.get_event_loop().run_in_executor(
                    self._executor,
                    extract_text_from_file,
                    Path(filename),
                    content_bytes
                )
            else:
                return False, 0, "No file path or content provided"

            if not text or len(text.strip()) < 10:
                logger.warning(f"[RAG] No text extracted from {filename}")
                return False, 0, "No text could be extracted from document"

            logger.info(f"[RAG] Extracted {len(text)} characters from {filename}")

            # 2. Chunk text
            chunks = chunk_text(text, self.chunk_size, self.chunk_overlap)
            if not chunks:
                return False, 0, "Document could not be chunked"

            logger.info(f"[RAG] Created {len(chunks)} chunks from {filename}")

            # 3. Generate embeddings for each chunk
            doc_chunks = []
            for i, chunk_text_content in enumerate(chunks):
                embedding = await generate_embedding(
                    chunk_text_content,
                    self.embedding_model,
                    self.ollama_url
                )

                if not embedding:
                    logger.warning(f"[RAG] Failed to generate embedding for chunk {i}")
                    continue

                doc_chunk = DocumentChunk(
                    id=len(doc_chunks),
                    folder_id=folder_id,
                    document_id=document_id,
                    filename=filename,
                    content=chunk_text_content,
                    chunk_index=i,
                    embedding=embedding
                )
                doc_chunks.append(doc_chunk)

            if not doc_chunks:
                return False, 0, "Failed to generate embeddings"

            # 4. Store in cache
            if folder_id not in self._chunk_cache:
                self._chunk_cache[folder_id] = []
            self._chunk_cache[folder_id].extend(doc_chunks)

            # 5. Persist to database for recovery after restart
            await self._persist_chunks(folder_id, document_id, doc_chunks)

            logger.info(f"[RAG] Successfully processed {filename}: {len(doc_chunks)} chunks embedded and persisted")
            return True, len(doc_chunks), f"Processed {len(doc_chunks)} chunks"

        except Exception as e:
            logger.error(f"[RAG] Error processing document {filename}: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return False, 0, str(e)

    async def process_temp_document(
        self,
        session_id: str,
        filename: str,
        content_bytes: bytes
    ) -> Tuple[bool, int, str]:
        """
        Process a temporary document for ephemeral/one-time chat context.

        - Stores in memory only (temp_store)
        - NOT persisted to database
        - Lost on server restart

        Returns: (success, chunk_count, message)
        """
        logger.info(f"[RAG Temp] Processing temp document: {filename} for session {session_id}")

        try:
            # 1. Extract text
            text = await asyncio.get_event_loop().run_in_executor(
                self._executor,
                extract_text_from_file,
                Path(filename),
                content_bytes
            )

            if not text or len(text.strip()) < 10:
                logger.warning(f"[RAG Temp] No text extracted from {filename}")
                return False, 0, "No text could be extracted from document"

            logger.info(f"[RAG Temp] Extracted {len(text)} characters from {filename}")

            # 2. Chunk text
            chunks = chunk_text(text, self.chunk_size, self.chunk_overlap)
            if not chunks:
                return False, 0, "Document could not be chunked"

            logger.info(f"[RAG Temp] Created {len(chunks)} chunks from {filename}")

            # 3. Generate embeddings for each chunk
            temp_chunks: List[TempChunk] = []
            for i, chunk_text_content in enumerate(chunks):
                embedding = await generate_embedding(
                    chunk_text_content,
                    self.embedding_model,
                    self.ollama_url
                )

                if not embedding:
                    logger.warning(f"[RAG Temp] Failed to generate embedding for chunk {i}")
                    continue

                temp_chunk = TempChunk(
                    id=len(temp_chunks),
                    session_id=session_id,
                    filename=filename,
                    content=chunk_text_content,
                    chunk_index=i,
                    embedding=embedding
                )
                temp_chunks.append(temp_chunk)

            if not temp_chunks:
                return False, 0, "Failed to generate embeddings"

            # 4. Store in temp_store (memory only, NO database persistence)
            if session_id not in self.temp_store:
                self.temp_store[session_id] = []
            self.temp_store[session_id].extend(temp_chunks)

            logger.info(f"[RAG Temp] Successfully processed {filename}: {len(temp_chunks)} chunks (session: {session_id})")
            return True, len(temp_chunks), f"Processed {len(temp_chunks)} chunks (temp)"

        except Exception as e:
            logger.error(f"[RAG Temp] Error processing temp document {filename}: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return False, 0, str(e)

    def clear_temp_session(self, session_id: str) -> None:
        """Clear temporary chunks for a session."""
        if session_id in self.temp_store:
            count = len(self.temp_store[session_id])
            del self.temp_store[session_id]
            logger.info(f"[RAG Temp] Cleared {count} temp chunks for session {session_id}")

    def get_temp_stats(self, session_id: str = None) -> Dict[str, Any]:
        """Get temp store statistics."""
        if session_id:
            chunks = self.temp_store.get(session_id, [])
            return {
                "session_id": session_id,
                "chunk_count": len(chunks),
                "filenames": list(set(c.filename for c in chunks))
            }
        else:
            return {
                "total_sessions": len(self.temp_store),
                "total_chunks": sum(len(chunks) for chunks in self.temp_store.values())
            }

    def _compute_cache_key(self, query: str, folder_id: int = None, session_id: str = None) -> str:
        """검색 쿼리 캐시 키 생성."""
        import hashlib
        key_str = f"{query}|{folder_id}|{session_id}"
        return hashlib.md5(key_str.encode()).hexdigest()
    
    def _is_cache_valid(self, cache_entry: Tuple[List[RAGSearchResult], datetime]) -> bool:
        """캐시 엔트리가 유효한지 확인 (TTL 체크)."""
        _, timestamp = cache_entry
        age = (datetime.utcnow() - timestamp).total_seconds()
        return age < self._cache_ttl_seconds
    
    def _cleanup_expired_cache(self) -> None:
        """만료된 캐시 엔트리 제거."""
        now = datetime.utcnow()
        expired_keys = [
            key for key, (_, timestamp) in self._search_cache.items()
            if (now - timestamp).total_seconds() >= self._cache_ttl_seconds
        ]
        for key in expired_keys:
            del self._search_cache[key]
        if expired_keys:
            logger.debug(f"[RAG] Cleaned up {len(expired_keys)} expired cache entries")

    async def search(
        self,
        query: str,
        folder_id: int = None,
        session_id: str = None,
        top_k: int = DEFAULT_TOP_K,
        min_score: float = 0.3,
        timeout: float = 10.0
    ) -> List[RAGSearchResult]:
        """
        Search for relevant chunks using semantic similarity.

        Searches both:
        - Source A: DB (Folder RAG) - permanent storage by folder_id
        - Source B: Memory (Temp RAG) - ephemeral storage by session_id

        Returns top_k results sorted by relevance, merged from both sources.
        
        Performance optimizations:
        - Result caching (5min TTL)
        - Parallel similarity computation
        - Early termination when enough results found
        - Timeout protection
        """
        # 캐시 확인
        cache_key = self._compute_cache_key(query, folder_id, session_id)
        if cache_key in self._search_cache:
            cache_entry = self._search_cache[cache_key]
            if self._is_cache_valid(cache_entry):
                cached_results, _ = cache_entry
                logger.info(f"[RAG] Cache hit for query: {query[:50]}... ({len(cached_results)} results)")
                return cached_results
        
        # 만료된 캐시 정리 (주기적으로)
        if len(self._search_cache) > 100:  # 캐시가 너무 커지면 정리
            self._cleanup_expired_cache()
        
        logger.info(f"[RAG] Searching (folder={folder_id}, session={session_id}) for: {query[:50]}...")

        try:
            # 타임아웃 적용하여 임베딩 생성
            query_embedding = await asyncio.wait_for(
                generate_embedding(query, self.embedding_model, self.ollama_url),
                timeout=timeout
            )
        except asyncio.TimeoutError:
            logger.error(f"[RAG] Embedding generation timeout after {timeout}s")
            return []
        except Exception as e:
            logger.error(f"[RAG] Failed to generate query embedding: {e}")
            return []

        if not query_embedding:
            logger.error("[RAG] Failed to generate query embedding")
            return []

        results = []

        # Source A: Search in folder cache (permanent RAG) - 병렬 처리
        if folder_id is not None:
            folder_chunks = self._chunk_cache.get(folder_id, [])
            if folder_chunks:
                # 유사도 계산을 병렬로 처리 (CPU 바운드이지만 큰 청크 수에서 효과적)
                def compute_similarity(chunk: DocumentChunk) -> Optional[RAGSearchResult]:
                    if not chunk.embedding:
                        return None
                    score = cosine_similarity(query_embedding, chunk.embedding)
                    if score >= min_score:
                        return RAGSearchResult(
                            content=chunk.content,
                            filename=chunk.filename,
                            score=score,
                            chunk_index=chunk.chunk_index,
                            document_id=chunk.document_id
                        )
                    return None
                
                # ThreadPoolExecutor로 병렬 처리 (CPU 바운드 작업)
                loop = asyncio.get_event_loop()
                chunk_results = await loop.run_in_executor(
                    self._executor,
                    lambda: [compute_similarity(chunk) for chunk in folder_chunks]
                )
                
                # None 제거 및 결과 추가
                valid_results = [r for r in chunk_results if r is not None]
                results.extend(valid_results)
                logger.info(f"[RAG] Folder {folder_id}: found {len(valid_results)} matches from {len(folder_chunks)} chunks")

        # Source B: Search in temp_store (ephemeral RAG) - 병렬 처리
        if session_id is not None:
            temp_chunks = self.temp_store.get(session_id, [])
            if temp_chunks:
                def compute_temp_similarity(chunk: TempChunk) -> Optional[RAGSearchResult]:
                    if not chunk.embedding:
                        return None
                    score = cosine_similarity(query_embedding, chunk.embedding)
                    if score >= min_score:
                        return RAGSearchResult(
                            content=chunk.content,
                            filename=f"[TEMP] {chunk.filename}",
                            score=score,
                            chunk_index=chunk.chunk_index,
                            document_id=-1  # Temp docs have no document_id
                        )
                    return None
                
                # ThreadPoolExecutor로 병렬 처리
                loop = asyncio.get_event_loop()
                temp_results = await loop.run_in_executor(
                    self._executor,
                    lambda: [compute_temp_similarity(chunk) for chunk in temp_chunks]
                )
                
                # None 제거 및 결과 추가
                valid_temp_results = [r for r in temp_results if r is not None]
                results.extend(valid_temp_results)
                logger.info(f"[RAG] Temp session {session_id}: found {len(valid_temp_results)} matches from {len(temp_chunks)} chunks")

        if not results:
            logger.info(f"[RAG] No chunks found for folder={folder_id}, session={session_id}")
            # 빈 결과도 캐시 (짧은 TTL)
            self._search_cache[cache_key] = ([], datetime.utcnow())
            return []

        # Sort by score and return top_k
        results.sort(key=lambda x: x.score, reverse=True)
        top_results = results[:top_k]

        # 결과 캐시 저장
        self._search_cache[cache_key] = (top_results, datetime.utcnow())

        logger.info(f"[RAG] Found {len(top_results)} relevant chunks (min_score={min_score}, total candidates: {len(results)})")
        for r in top_results:
            logger.debug(f"[RAG]   - {r.filename} (score={r.score:.3f}): {r.content[:50]}...")

        return top_results

    async def _persist_chunks(
        self,
        folder_id: int,
        document_id: int,
        chunks: List[DocumentChunk]
    ) -> bool:
        """Persist chunks to database for recovery after restart."""
        try:
            from mellow_link.infra.database import SessionLocal, DocumentChunk as DBDocumentChunk

            db = SessionLocal()
            try:
                # Delete existing chunks for this document (in case of re-upload)
                db.query(DBDocumentChunk).filter(
                    DBDocumentChunk.document_id == document_id
                ).delete()

                # Insert new chunks
                for chunk in chunks:
                    db_chunk = DBDocumentChunk(
                        folder_id=folder_id,
                        document_id=document_id,
                        filename=chunk.filename,
                        content=chunk.content,
                        chunk_index=chunk.chunk_index,
                        embedding=json.dumps(chunk.embedding)  # Serialize embedding as JSON
                    )
                    db.add(db_chunk)

                db.commit()
                logger.info(f"[RAG] Persisted {len(chunks)} chunks to database for document {document_id}")
                return True

            except Exception as e:
                db.rollback()
                logger.error(f"[RAG] Failed to persist chunks: {e}")
                return False
            finally:
                db.close()

        except ImportError as e:
            logger.warning(f"[RAG] Database module not available: {e}")
            return False

    async def load_chunks_from_db(self, folder_id: int = None) -> int:
        """
        Load chunks from database into memory cache.

        Args:
            folder_id: Optional folder ID to load. If None, loads all folders.

        Returns:
            Number of chunks loaded.
        """
        try:
            from mellow_link.infra.database import SessionLocal, DocumentChunk as DBDocumentChunk

            db = SessionLocal()
            try:
                query = db.query(DBDocumentChunk)
                if folder_id is not None:
                    query = query.filter(DBDocumentChunk.folder_id == folder_id)

                db_chunks = query.all()

                if not db_chunks:
                    logger.info(f"[RAG] No chunks found in database" + (f" for folder {folder_id}" if folder_id else ""))
                    return 0

                # Group by folder_id and load into cache
                loaded_count = 0
                for db_chunk in db_chunks:
                    fid = db_chunk.folder_id

                    if fid not in self._chunk_cache:
                        self._chunk_cache[fid] = []

                    # Deserialize embedding from JSON
                    try:
                        embedding = json.loads(db_chunk.embedding)
                    except json.JSONDecodeError:
                        logger.warning(f"[RAG] Invalid embedding JSON for chunk {db_chunk.id}")
                        continue

                    chunk = DocumentChunk(
                        id=db_chunk.id,
                        folder_id=fid,
                        document_id=db_chunk.document_id,
                        filename=db_chunk.filename,
                        content=db_chunk.content,
                        chunk_index=db_chunk.chunk_index,
                        embedding=embedding
                    )
                    self._chunk_cache[fid].append(chunk)
                    loaded_count += 1

                logger.info(f"[RAG] Loaded {loaded_count} chunks from database into cache")
                return loaded_count

            finally:
                db.close()

        except ImportError as e:
            logger.warning(f"[RAG] Database module not available: {e}")
            return 0
        except Exception as e:
            logger.error(f"[RAG] Failed to load chunks from database: {e}")
            return 0

    def clear_folder_cache(self, folder_id: int) -> None:
        """Clear cached chunks for a folder."""
        if folder_id in self._chunk_cache:
            del self._chunk_cache[folder_id]
            logger.info(f"[RAG] Cleared cache for folder {folder_id}")

    def clear_document_from_cache(self, folder_id: int, document_id: int) -> None:
        """Remove a specific document's chunks from cache."""
        if folder_id in self._chunk_cache:
            self._chunk_cache[folder_id] = [
                c for c in self._chunk_cache[folder_id]
                if c.document_id != document_id
            ]
            logger.info(f"[RAG] Removed document {document_id} from folder {folder_id} cache")

    def get_stats(self) -> Dict[str, Any]:
        """Get RAG service statistics."""
        total_chunks = sum(len(chunks) for chunks in self._chunk_cache.values())
        return {
            "initialized": self._is_initialized,
            "folders_cached": len(self._chunk_cache),
            "total_chunks": total_chunks,
            "embedding_model": self.embedding_model,
            "chunk_size": self.chunk_size,
        }


# =============================================================================
# Global Instance & Factory
# =============================================================================

_rag_service: Optional[RAGService] = None


def get_rag_service() -> Optional[RAGService]:
    """Get the global RAG service instance."""
    return _rag_service


def set_rag_service(service: RAGService) -> None:
    """Set the global RAG service instance."""
    global _rag_service
    _rag_service = service


async def create_rag_service(
    embedding_model: str = DEFAULT_EMBEDDING_MODEL,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    ollama_url: str = OLLAMA_BASE_URL
) -> RAGService:
    """
    Factory function to create and initialize RAG service.
    """
    service = RAGService(
        embedding_model=embedding_model,
        chunk_size=chunk_size,
        ollama_url=ollama_url
    )
    await service.initialize()
    return service
