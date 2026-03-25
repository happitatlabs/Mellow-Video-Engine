"""
Document Service - CPU-based Document Generation

This module provides document generation capabilities for Word documents.
Runs ENTIRELY on CPU in a thread pool to avoid blocking the async event loop.

CRITICAL: All CPU-bound operations MUST use loop.run_in_executor() to prevent
blocking the async event loop while GPU tasks (LLM/Image) are running.

Libraries:
    - Word: python-docx
    - PDF: reportlab (optional - basic implementation)

Design:
    - All CPU-bound operations run in ThreadPoolExecutor
    - Non-blocking async interface
    - Does NOT touch the GPU
"""

import asyncio
import logging
import time
from concurrent.futures import ThreadPoolExecutor
from typing import Optional, Dict, Any, List
from dataclasses import dataclass, field
from enum import Enum, auto
from pathlib import Path
from datetime import datetime

logger = logging.getLogger(__name__)


# =============================================================================
# Enums and Data Classes
# =============================================================================

class DocumentType(Enum):
    """Supported document output types."""
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    MARKDOWN = "md"


class DocumentStatus(Enum):
    """Document service status."""
    IDLE = auto()
    PROCESSING = auto()
    ERROR = auto()


@dataclass
class DocumentRequest:
    """
    Request structure for document generation.

    Attributes:
        content: Main content (text or Markdown)
        output_type: Desired output format (default: DOCX)
        title: Document title
        metadata: Additional metadata (author, date, etc.)
        style_options: Formatting options (fonts, margins, etc.)
        filename: Optional custom filename
    """
    content: str
    output_type: DocumentType = DocumentType.DOCX
    title: str = "Untitled"
    metadata: Dict[str, str] = field(default_factory=dict)
    style_options: Dict[str, Any] = field(default_factory=dict)
    filename: Optional[str] = None


@dataclass
class DocumentResult:
    """
    Result structure from document generation.

    Attributes:
        output_path: Path to generated document
        output_type: Type of document generated
        page_count: Estimated number of pages
        file_size_bytes: Size of generated file
        generation_time_ms: Time to generate
    """
    output_path: Path
    output_type: DocumentType
    page_count: int = 0
    file_size_bytes: int = 0
    generation_time_ms: float = 0.0


class DocumentGenerationError(Exception):
    """Exception for document generation failures."""
    pass


# =============================================================================
# Document Service Class
# =============================================================================

class DocumentService:
    """
    Service for CPU-based document generation.

    CRITICAL DESIGN:
        All CPU-intensive operations run in a ThreadPoolExecutor via
        loop.run_in_executor(). This ensures the async event loop is
        NEVER blocked, allowing GPU tasks to continue uninterrupted.

    This service:
        - Generates Word documents (python-docx)
        - Generates basic PDF documents (reportlab)
        - Generates HTML from Markdown
        - Does NOT use the GPU

    Usage:
        service = DocumentService()
        await service.initialize()

        request = DocumentRequest(content="Hello World", title="Test")
        result = await service.generate(request)

        await service.shutdown()
    """

    DEFAULT_OUTPUT_DIR: Path = Path("./outputs/documents")
    DEFAULT_THREAD_WORKERS: int = 2  # Keep low - CPU only

    def __init__(
        self,
        output_dir: Optional[Path] = None,
        max_workers: int = DEFAULT_THREAD_WORKERS
    ):
        """
        Initialize Document Service.

        Args:
            output_dir: Directory for generated documents
            max_workers: Maximum thread pool workers (default: 2)
        """
        self.output_dir = output_dir or self.DEFAULT_OUTPUT_DIR
        self.max_workers = max_workers

        self._status: DocumentStatus = DocumentStatus.IDLE
        self._executor: Optional[ThreadPoolExecutor] = None
        self._is_initialized: bool = False

    # ==================== Lifecycle ====================

    async def initialize(self) -> None:
        """
        Initialize the document service.

        Creates:
            - Output directory
            - Thread pool for CPU-bound work
        """
        logger.info("[DocumentService] Initializing...")

        # Create output directory
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Initialize thread pool for CPU-bound operations
        self._executor = ThreadPoolExecutor(
            max_workers=self.max_workers,
            thread_name_prefix="doc_worker"
        )

        self._is_initialized = True
        self._status = DocumentStatus.IDLE
        logger.info(f"[DocumentService] Initialized with {self.max_workers} workers")

    async def shutdown(self) -> None:
        """Shutdown the document service and cleanup thread pool."""
        logger.info("[DocumentService] Shutting down...")

        if self._executor:
            self._executor.shutdown(wait=True)
            self._executor = None

        self._is_initialized = False
        logger.info("[DocumentService] Shutdown complete")

    async def connect(self) -> bool:
        """Connect alias for orchestrator compatibility."""
        await self.initialize()
        return True

    async def disconnect(self) -> None:
        """Disconnect alias for orchestrator compatibility."""
        await self.shutdown()

    def get_status(self) -> DocumentStatus:
        """Get current service status."""
        return self._status

    def is_ready(self) -> bool:
        """Check if service is ready."""
        return self._is_initialized and self._status == DocumentStatus.IDLE

    def is_available(self) -> bool:
        """Check if service is available (orchestrator compatibility)."""
        return self._is_initialized

    # ==================== Document Generation ====================

    async def generate(self, request: DocumentRequest) -> DocumentResult:
        """
        Generate a document from request.

        CRITICAL: This method delegates CPU work to run_in_executor()
        to avoid blocking the async event loop.

        Args:
            request: DocumentRequest with content and options

        Returns:
            DocumentResult with generated file path

        Raises:
            DocumentGenerationError: If generation fails
        """
        if not self._is_initialized:
            raise DocumentGenerationError("DocumentService not initialized")

        self._status = DocumentStatus.PROCESSING
        start_time = time.time()

        try:
            # Generate output path
            output_path = self._get_output_path(
                request.filename or request.title,
                request.output_type
            )

            # Generate based on type - ALL use run_in_executor
            if request.output_type == DocumentType.DOCX:
                result = await self._generate_docx(
                    request.content,
                    output_path,
                    request.title,
                    request.style_options
                )
            elif request.output_type == DocumentType.PDF:
                result = await self._generate_pdf(
                    request.content,
                    output_path,
                    request.title,
                    request.style_options
                )
            elif request.output_type == DocumentType.HTML:
                result = await self._generate_html(
                    request.content,
                    output_path,
                    request.title
                )
            else:
                # Markdown - simple write
                result = await self._generate_markdown(
                    request.content,
                    output_path
                )

            result.generation_time_ms = (time.time() - start_time) * 1000

            logger.info(
                f"[DocumentService] Generated {request.output_type.value}: "
                f"{output_path.name} ({result.generation_time_ms:.0f}ms)"
            )
            return result

        except Exception as e:
            logger.error(f"[DocumentService] Generation failed: {e}")
            raise DocumentGenerationError(str(e))

        finally:
            self._status = DocumentStatus.IDLE

    async def _generate_docx(
        self,
        content: str,
        output_path: Path,
        title: str,
        options: Dict[str, Any]
    ) -> DocumentResult:
        """
        Generate Word document using python-docx.

        CRITICAL: Uses run_in_executor to run in thread pool.
        """
        loop = asyncio.get_event_loop()

        def _create_docx() -> int:
            """CPU-bound DOCX creation - runs in thread pool."""
            try:
                from docx import Document
                from docx.shared import Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError:
                raise DocumentGenerationError(
                    "python-docx not installed. Run: pip install python-docx"
                )

            doc = Document()

            # Add title
            doc.add_heading(title, level=0)

            # Configure style
            style = doc.styles['Normal']
            font = style.font
            font.name = options.get("font_name", "Calibri")
            font.size = Pt(options.get("font_size", 11))

            # Add content paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            doc.save(str(output_path))

            # Estimate page count
            return max(1, len(content) // 3000)

        # CRITICAL: Run in executor to avoid blocking async loop
        page_count = await loop.run_in_executor(self._executor, _create_docx)

        return DocumentResult(
            output_path=output_path,
            output_type=DocumentType.DOCX,
            page_count=page_count,
            file_size_bytes=output_path.stat().st_size
        )

    async def _generate_pdf(
        self,
        content: str,
        output_path: Path,
        title: str,
        options: Dict[str, Any]
    ) -> DocumentResult:
        """
        Generate PDF document using reportlab.

        CRITICAL: Uses run_in_executor to run in thread pool.
        """
        loop = asyncio.get_event_loop()

        def _create_pdf() -> int:
            """CPU-bound PDF creation - runs in thread pool."""
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import cm
                from reportlab.lib.enums import TA_JUSTIFY
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            except ImportError:
                raise DocumentGenerationError(
                    "reportlab not installed. Run: pip install reportlab"
                )

            # Create document
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=A4,
                leftMargin=2.5 * cm,
                rightMargin=2.5 * cm,
                topMargin=2.5 * cm,
                bottomMargin=2.5 * cm
            )

            # Styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30
            )
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontSize=options.get("font_size", 11),
                alignment=TA_JUSTIFY,
                spaceBefore=6,
                spaceAfter=6,
                leading=14
            )

            # Build content
            story = []
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 12))

            # Add paragraphs
            for para in content.split('\n\n'):
                if para.strip():
                    para_text = para.replace('\n', '<br/>')
                    story.append(Paragraph(para_text, body_style))

            doc.build(story)

            # Estimate page count
            return max(1, output_path.stat().st_size // 3000)

        # CRITICAL: Run in executor to avoid blocking async loop
        page_count = await loop.run_in_executor(self._executor, _create_pdf)

        return DocumentResult(
            output_path=output_path,
            output_type=DocumentType.PDF,
            page_count=page_count,
            file_size_bytes=output_path.stat().st_size
        )

    async def _generate_html(
        self,
        content: str,
        output_path: Path,
        title: str
    ) -> DocumentResult:
        """
        Generate HTML document.

        CRITICAL: Uses run_in_executor to run in thread pool.
        """
        loop = asyncio.get_event_loop()

        def _create_html():
            """CPU-bound HTML creation - runs in thread pool."""
            # Try to convert markdown if available
            html_content = content
            try:
                import markdown
                html_content = markdown.markdown(
                    content,
                    extensions=['tables', 'fenced_code', 'nl2br']
                )
            except ImportError:
                # Basic fallback
                html_content = content.replace('\n\n', '</p><p>')
                html_content = f"<p>{html_content}</p>"

            html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
        }}
        h1 {{ color: #333; }}
        p {{ text-align: justify; }}
        code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; }}
        pre {{ background: #f4f4f4; padding: 15px; overflow-x: auto; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    {html_content}
</body>
</html>"""
            output_path.write_text(html, encoding='utf-8')

        # CRITICAL: Run in executor to avoid blocking async loop
        await loop.run_in_executor(self._executor, _create_html)

        return DocumentResult(
            output_path=output_path,
            output_type=DocumentType.HTML,
            page_count=1,
            file_size_bytes=output_path.stat().st_size
        )

    async def _generate_markdown(
        self,
        content: str,
        output_path: Path
    ) -> DocumentResult:
        """
        Save content as Markdown file.

        Even simple I/O uses run_in_executor for consistency.
        """
        loop = asyncio.get_event_loop()

        def _write_file():
            output_path.write_text(content, encoding='utf-8')

        # CRITICAL: Run in executor to avoid blocking async loop
        await loop.run_in_executor(self._executor, _write_file)

        return DocumentResult(
            output_path=output_path,
            output_type=DocumentType.MARKDOWN,
            page_count=1,
            file_size_bytes=output_path.stat().st_size
        )

    # ==================== Orchestrator Interface ====================

    async def execute(self, request_data: Dict[str, Any]) -> DocumentResult:
        """
        Execute method for orchestrator compatibility.

        Args:
            request_data: Dict with generation parameters

        Returns:
            DocumentResult
        """
        doc_type_str = request_data.get("output_type", "docx").lower()
        doc_type_map = {
            "pdf": DocumentType.PDF,
            "docx": DocumentType.DOCX,
            "html": DocumentType.HTML,
            "md": DocumentType.MARKDOWN,
            "markdown": DocumentType.MARKDOWN,
        }

        request = DocumentRequest(
            content=request_data.get("content", ""),
            output_type=doc_type_map.get(doc_type_str, DocumentType.DOCX),
            title=request_data.get("title", "Document"),
            metadata=request_data.get("metadata", {}),
            style_options=request_data.get("style_options", {}),
            filename=request_data.get("filename"),
        )
        return await self.generate(request)

    # ==================== Utilities ====================

    def _get_output_path(self, filename: str, doc_type: DocumentType) -> Path:
        """Generate unique output path for document."""
        # Clean filename
        safe_name = "".join(c for c in filename if c.isalnum() or c in (' ', '-', '_'))
        safe_name = safe_name.strip() or "document"

        # Add timestamp for uniqueness
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        return self.output_dir / f"{safe_name}_{timestamp}.{doc_type.value}"

    async def health_check(self) -> Dict[str, Any]:
        """Perform health check."""
        return {
            "healthy": self._is_initialized,
            "status": self._status.name,
            "output_dir": str(self.output_dir),
            "max_workers": self.max_workers,
        }

    def to_dict(self) -> Dict[str, Any]:
        """Export service state as dictionary."""
        return {
            "status": self._status.name,
            "is_initialized": self._is_initialized,
            "output_dir": str(self.output_dir),
            "max_workers": self.max_workers,
        }


# =============================================================================
# Factory Function
# =============================================================================

def create_document_service(
    output_dir: Optional[Path] = None,
    max_workers: int = 2
) -> DocumentService:
    """
    Factory function to create DocumentService.

    Args:
        output_dir: Output directory for documents
        max_workers: Thread pool size (default: 2)

    Returns:
        Configured DocumentService instance
    """
    return DocumentService(
        output_dir=output_dir,
        max_workers=max_workers
    )
